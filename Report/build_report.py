"""
Build the CS516 project report as a .docx file.
Pulls screenshots from ../Screenshots and source code excerpts from ../Python and ../Cpp.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SHOTS = os.path.join(ROOT, "Screenshots")
OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "CS516_Project_Report.docx")

doc = Document()

# ---------- page setup ----------
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_run(p, text, *, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if font:
        r.font.name = font
    return r


def H(text, level=1, after=6):
    """Add a heading-ish line with our own formatting (not built-in headings,
    so the layout matches the sample report style)."""
    p = doc.add_paragraph()
    sizes = {1: 16, 2: 13, 3: 12}
    add_run(p, text, bold=True, size=sizes.get(level, 12), color="1F3A5F")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    return p


def P(text, *, bold=False, italic=False, size=11, indent=False, align=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, text, bold=bold, italic=italic, size=size)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    add_run(p, text)
    return p


def code_block(text, *, mono="Consolas", size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, font=mono, size=size)
    # shaded background by setting paragraph border + shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F6F8")
    pPr.append(shd)
    return p


def figure(image_name, caption, width_in=6.0):
    path = os.path.join(SHOTS, image_name)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, italic=True, size=9, color="5F6B7C")
    cap.paragraph_format.space_after = Pt(8)


# =================================================================
#                          TITLE PAGE
# =================================================================
title_block = doc.add_paragraph()
title_block.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(title_block, "\nIMAM ABDULRAHMAN BIN FAISAL UNIVERSITY\n", bold=True, size=14, color="1F3A5F")
add_run(title_block, "College of Computer Science and Information Technology\n", size=11, color="3A4A5F")
add_run(title_block, "\n\n", size=11)
add_run(title_block, "CS 516 – Advanced Programming Language\n", bold=True, size=18, color="1F3A5F")
add_run(title_block, "2nd Semester 2025/2026 — Term Project\n", size=12, color="5F6B7C")
add_run(title_block, "\n\n", size=11)
add_run(title_block, "Smart Student Management\n& Attendance System\n",
        bold=True, size=22, color="1F3A5F")
add_run(title_block, "\nA Client–Server Application implemented in Python and C++\n",
        italic=True, size=12, color="5F6B7C")

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Report Submitted to:\n", bold=True, size=12)
add_run(p, "Dr. Dhiaa, CCSIT-IAU\n", size=12, color="3A4A5F")
doc.add_paragraph()

# members table
p = doc.add_paragraph()
add_run(p, "Submitted by:", bold=True, size=12)

members = [
    ("Ahmed Mabkhot Al-Awlaqi", "2220007240", "Team Leader"),
    ("Abdulaziz Alghamdi",       "2220003178", "Member"),
    ("Abdullah Alsaedi",         "2220002306", "Member"),
]
tbl = doc.add_table(rows=1 + len(members), cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, h in enumerate(["Name", "Student ID", "Role"]):
    hdr[i].text = ""
    add_run(hdr[i].paragraphs[0], h, bold=True)
for i, (n, sid, role) in enumerate(members, start=1):
    tbl.rows[i].cells[0].text = n
    tbl.rows[i].cells[1].text = sid
    tbl.rows[i].cells[2].text = role

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Declaration: ", bold=True, italic=True, size=10, color="5F6B7C")
add_run(p,
        "We hereby declare that this report and the accompanying source code are our own original work. "
        "Any external references that helped us along the way (documentation, tutorials, lecture material) "
        "are acknowledged at the end.",
        italic=True, size=10, color="5F6B7C")

doc.add_page_break()

# =================================================================
#                        TABLE OF CONTENTS
# =================================================================
H("Table of Contents", level=1)
toc_items = [
    ("1. Problem Statement", 3),
    ("2. Project Goals and Objectives / Deliverables", 3),
    ("3. Introduction", 4),
    ("4. Project Scope", 4),
    ("5. Success Factors and Benefits", 5),
    ("6. Limitations / Restrictions", 5),
    ("7. Selected Programming Language 1: Python", 6),
    ("    A. Login Interface", 6),
    ("    B. Students Management", 7),
    ("    C. Courses Management", 9),
    ("    D. Enrollment", 10),
    ("    E. Attendance", 11),
    ("8. Selected Programming Language 2: C++", 12),
    ("    A. Console Login & Main Menu", 12),
    ("    B. Students Management", 13),
    ("    C. Courses Management", 15),
    ("    D. Enrollment", 15),
    ("    E. Attendance", 16),
    ("9. Comparison Between the Two Programming Languages", 17),
    ("    Part 1: Syntax, Dependencies, Paradigm, Memory Management", 17),
    ("    Part 2: Side-by-Side Function Comparison", 19),
    ("10. Attachments", 23),
    ("11. References", 23),
]
for title, page in toc_items:
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=2)  # dotted leader
    add_run(p, title, size=11)
    add_run(p, f"\t{page}", size=11)

doc.add_page_break()

# Table of Figures
H("Table of Figures", level=1)
figures_list = [
    "Figure 1. Login window of the Python client",
    "Figure 2. Main window — Students tab (admin view)",
    "Figure 3. Add Student form filled in (validation about to fire)",
    "Figure 4. Students list after a new student has been added",
    "Figure 5. Searching the students list by name",
    "Figure 6. Courses tab — default state",
    "Figure 7. Courses list after adding CS401 Database Systems",
    "Figure 8. Enrollment tab listing all current enrollments",
    "Figure 9. Attendance report (full)",
    "Figure 10. Per-course attendance summary",
    "Figure 11. C++ client — login prompt",
    "Figure 12. C++ client — main menu",
    "Figure 13. C++ client — Students view-all",
    "Figure 14. C++ client — Add student rejected (bad ID)",
    "Figure 15. C++ client — Add student success",
    "Figure 16. C++ client — Courses view-all",
    "Figure 17. C++ client — Enrollment list",
    "Figure 18. C++ client — Enrolling a student",
    "Figure 19. C++ client — Attendance report for CS516",
    "Figure 20. C++ client — Attendance summary for CS516",
    "Figure 21. C++ client — Bad date format rejected",
]
for i, f in enumerate(figures_list, start=1):
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
    add_run(p, f, size=10)
    add_run(p, f"\t{6 + i // 2}", size=10)

doc.add_page_break()

# Table of Tables
H("Table of Tables", level=1)
tables_list = [
    ("Table 1. Language comparison — syntax / dependencies / paradigm / memory management", 17),
    ("Table 2. Default user accounts used for demo", 6),
    ("Table 3. Data model summary", 6),
]
for t, pg in tables_list:
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=2)
    add_run(p, t, size=10)
    add_run(p, f"\t{pg}", size=10)

doc.add_page_break()

# =================================================================
#                       1. PROBLEM STATEMENT
# =================================================================
H("1. Problem Statement", level=1)
P(
    "Most academic departments still juggle a mix of paper attendance sheets, Excel files, and "
    "the occasional WhatsApp group to keep track of who showed up to class. That works until it "
    "doesn’t — sheets get lost, two instructors mark the same student twice, and at the end of "
    "the semester somebody has to manually count rows to figure out who exceeded the absence limit. "
    "Our project starts from that everyday frustration: a department needs one simple place to keep "
    "students, courses, enrollments, and daily attendance, and it needs to be quick enough that an "
    "instructor can mark a class in under a minute.",
    align="justify",
)
P(
    "We chose a client–server design on purpose. The data lives in one place (the server) and any "
    "number of clients on the network can read or update it without stepping on each other. "
    "Building the same system twice — once in Python with a desktop GUI and once in C++ as a "
    "console program — also lets us see, in practice, how the two paradigms differ when you "
    "actually have to ship a working system, not just a textbook snippet.",
    align="justify",
)

# =================================================================
#                       2. GOALS / DELIVERABLES
# =================================================================
H("2. Project Goals and Objectives / Deliverables", level=1)
H("A. Project Goals", level=2)
for b in [
    "Deliver a working client–server system that handles students, courses, enrollment and attendance.",
    "Implement the same feature set in two languages from different paradigms (Python, C++).",
    "Validate every user input on the server side so the database never holds garbage.",
    "Make the GUI usable for someone who has never seen the system before — no manual required.",
    "Compare the two implementations honestly, including where one beat the other.",
]:
    bullet(b)

H("B. Deliverables", level=2)
for b in [
    "A Python implementation: tkinter GUI client + TCP socket server + SQLite database.",
    "A C++ implementation: console client + TCP socket server + flat-file storage.",
    "A written report describing both systems and comparing them.",
    "All source code, ready to compile and run, attached at the end of this report.",
]:
    bullet(b)

# =================================================================
#                         3. INTRODUCTION
# =================================================================
H("3. Introduction", level=1)
P(
    "The Smart Student Management & Attendance System (we usually just call it Smart SMS) is a "
    "small client–server application aimed at department administrators and course instructors. "
    "The administrator uses it to register new students, create courses, enroll students into "
    "the courses they’ll take, and so on. The instructor uses it during or after a lecture to "
    "mark who was present and who wasn’t, and to pull a quick summary at any point in the term.",
    align="justify",
)
P(
    "Because we wanted to learn something about each paradigm rather than just hit the rubric, we "
    "built two complete versions. The Python one leans on its standard library — tkinter for the "
    "UI, sqlite3 for storage, and sockets for the client/server channel — so a single command "
    "(python server.py) is enough to run the backend on any machine that already has Python. "
    "The C++ one targets the same protocol but uses object-oriented design with a Repository "
    "per entity, the Winsock API for networking, and plain TAB-separated files for storage. "
    "Both halves talk to each other through the same simple line-based protocol, which means a "
    "Python client could actually connect to the C++ server if you wanted it to — although in "
    "this report we keep them paired.",
    align="justify",
)

# Default accounts table
H("Default accounts (for demo only)", level=3)
acc_tbl = doc.add_table(rows=4, cols=3)
acc_tbl.style = "Light Grid Accent 1"
acc_tbl.rows[0].cells[0].text = "Username"
acc_tbl.rows[0].cells[1].text = "Password"
acc_tbl.rows[0].cells[2].text = "Role"
for i, row in enumerate([("admin", "admin123", "admin"),
                          ("ahmed", "ahmed123", "instructor"),
                          ("abdullah", "abdullah123", "instructor")], start=1):
    for j, v in enumerate(row):
        acc_tbl.rows[i].cells[j].text = v
P("Table 2. Default user accounts used for demo.", italic=True, size=9)

H("Data model at a glance", level=3)
dm_tbl = doc.add_table(rows=6, cols=2)
dm_tbl.style = "Light Grid Accent 1"
rows = [
    ("Table", "Fields"),
    ("users", "username, password, role"),
    ("students", "student_id, name, email, phone, major, level"),
    ("courses", "code, title, credits, instructor"),
    ("enrollments", "student_id, course_code, semester"),
    ("attendance", "student_id, course_code, adate, status"),
]
for i, (a, b) in enumerate(rows):
    dm_tbl.rows[i].cells[0].text = a
    dm_tbl.rows[i].cells[1].text = b
    if i == 0:
        for c in dm_tbl.rows[i].cells:
            for r in c.paragraphs[0].runs:
                r.bold = True
P("Table 3. Data model summary.", italic=True, size=9)

# =================================================================
#                       4. PROJECT SCOPE
# =================================================================
H("4. Project Scope", level=1)
P(
    "Smart SMS is built for a single academic department, not a whole university. "
    "It assumes a small set of users (one or two admins plus a handful of instructors) and a "
    "manageable number of students — basically anything that fits comfortably in SQLite "
    "or a small text file. The features we cover are:",
    align="justify",
)
for b in [
    "User login with two roles (admin, instructor).",
    "Add / update / delete / search students (admin only) and view (everyone).",
    "Add / update / delete / search courses (admin only) and view (everyone).",
    "Enroll and unenroll students from courses for a given semester.",
    "Mark attendance per student/course/date (Present or Absent).",
    "Pull a full attendance report or a per-course summary at any time.",
]:
    bullet(b)
P(
    "Anything beyond that — grades, transcripts, fees, multi-campus support, mobile apps — is "
    "out of scope for this project. We deliberately kept the surface area small so we could "
    "implement it twice, properly, instead of once half-way.",
    align="justify",
)

# =================================================================
#                  5. SUCCESS FACTORS AND BENEFITS
# =================================================================
H("5. Success Factors and Benefits", level=1)
P(
    "The system is considered successful when an admin can register a student, enroll them in "
    "a course, and have an instructor mark attendance on that student within the same minute, "
    "across both implementations, without losing data or hitting any unhandled error. "
    "Concretely the benefits are:",
    align="justify",
)
for b in [
    "One source of truth: the server owns the data, the clients are thin.",
    "Validation lives in one place (server side), which means both clients get it for free.",
    "Attendance summaries that used to take ten minutes of counting rows now take one click.",
    "Because the protocol is plain text, debugging a connection problem is as easy as opening telnet.",
]:
    bullet(b)

# =================================================================
#                     6. LIMITATIONS / RESTRICTIONS
# =================================================================
H("6. Limitations / Restrictions", level=1)
for b in [
    "Passwords are stored in plain text. This is fine for a course project but would obviously "
    "need hashing (bcrypt / argon2) before any real deployment.",
    "The C++ server uses simple flat files. It is concurrency-safe within a single process thanks "
    "to a mutex, but it is not a real database — there are no transactions across files.",
    "Only one admin role and one instructor role are supported. There is no hierarchy of "
    "permissions beyond that split.",
    "The clients connect over plain TCP on localhost. There is no TLS, so don’t run this over the "
    "open internet without first putting it behind a tunnel or wrapping the socket.",
    "The system runs in English only; we did not add Arabic localisation.",
]:
    bullet(b)

doc.add_page_break()

# =================================================================
#                  7. SELECTED LANGUAGE 1: PYTHON
# =================================================================
H("7. Selected Programming Language 1: Python", level=1)
P(
    "Python is the friendlier of the two languages we picked. It’s multi-paradigm — you can write "
    "it as procedural code, as object-oriented code, or in a more functional style — and its "
    "standard library already ships with everything we needed: tkinter for the GUI, socket and "
    "threading for the server, and sqlite3 for storage. No external dependency manager, no build "
    "step. To run the system, you start the server in one terminal and the client in another:",
    align="justify",
)
code_block("$ python server.py    # in the Python/ folder, terminal #1\n"
           "$ python client.py    # terminal #2 — opens the GUI")

H("A. Login Interface", level=2)
P(
    "The first window the user sees is the login. Username and password go in, the client opens "
    "a TCP socket to the server, sends a JSON request {action: \"login\", payload: {...}}, and "
    "either gets back the user’s role or an \"Invalid credentials\" error. The window is "
    "deliberately tiny — there’s no point making people scroll just to type two fields.",
    align="justify",
)
figure("py_01_login.png", "Figure 1. Login window of the Python client.", width_in=4.5)

H("B. Students Management", level=2)
P(
    "Once you’re in, the Students tab is the default view. The top half is a form for the six "
    "student fields, and the bottom half is a tree-view that always shows the current list. "
    "The buttons in between are colour-coded so you can’t miss them: green for Add, blue for "
    "Update, red for Delete.",
    align="justify",
)
figure("py_02_main_students.png", "Figure 2. Main window — Students tab (admin view).")
P(
    "Clicking a row copies its values into the form, which makes editing painless. If you try to "
    "add a student with an ID that isn’t exactly ten digits, or an email that doesn’t contain @, "
    "the server rejects the request and the client pops up a clear message. Below is the form "
    "filled with intentionally broken values to show what the validation looks like in practice:",
    align="justify",
)
figure("py_03_students_form_filled.png", "Figure 3. Add Student form filled in (validation about to fire).")
P(
    "When the inputs are valid, the record is inserted and the list refreshes immediately. We "
    "added a fourth student (a Demo Student placeholder) below; notice that the sort order is preserved and "
    "the form clears so the next entry doesn’t accidentally inherit the old values:",
    align="justify",
)
figure("py_04_students_after_add.png", "Figure 4. Students list after a new student has been added.")
P(
    "Searching uses whatever field you type into — if you put a name in the Name box, it matches "
    "by name; an ID, by ID; a major, by major. The query goes to the server as a LIKE %keyword%, "
    "so partial matches work too:",
    align="justify",
)
figure("py_05_students_search.png", "Figure 5. Searching the students list by name.")

H("C. Courses Management", level=2)
P(
    "Courses follow exactly the same pattern as students — same form-on-top, list-on-bottom "
    "layout, same set of buttons. The course code is the primary key and is forced to upper case "
    "on the server so people can’t accidentally create both \"cs516\" and \"CS516\".",
    align="justify",
)
figure("py_06_courses_tab.png", "Figure 6. Courses tab — default state.")
P("Here is the list after adding CS401 Database Systems through the GUI:")
figure("py_07_courses_after_add.png", "Figure 7. Courses list after adding CS401 Database Systems.")

H("D. Enrollment", level=2)
P(
    "Enrollment is the bridge between students and courses. Pick a student ID, pick a course "
    "code, type the semester, hit Enroll. The server checks that both the student and the course "
    "exist before it accepts the row, and it prevents you from enrolling the same student into "
    "the same course twice in the same semester (the UNIQUE constraint takes care of that).",
    align="justify",
)
figure("py_08_enrollment_tab.png", "Figure 8. Enrollment tab listing all current enrollments.")

H("E. Attendance", level=2)
P(
    "The Attendance tab has two main jobs: marking the daily attendance, and producing reports. "
    "To mark, you pick a student, the course they’re enrolled in, the date, and choose Present "
    "or Absent from the dropdown. The system refuses to record attendance for a student who "
    "isn’t enrolled in that course — which prevents the classic copy-paste mistake where "
    "you mark someone present for the wrong class.",
    align="justify",
)
figure("py_09_attendance_report.png", "Figure 9. Attendance report (full).")
P(
    "The Summary by Course button is the one instructors actually care about. It groups by "
    "student and counts how many Presents versus Absents each one has for the course you "
    "specified — so figuring out who’s over the absence limit takes one glance:",
    align="justify",
)
figure("py_10_attendance_summary.png", "Figure 10. Per-course attendance summary.")

doc.add_page_break()

# =================================================================
#                   8. SELECTED LANGUAGE 2: C++
# =================================================================
H("8. Selected Programming Language 2: C++", level=1)
P(
    "C++ is the heavier sibling. It compiles to a native binary, has no garbage collector, and "
    "you’re responsible for picking the right data structures yourself. We chose it as our "
    "second language precisely because it sits on the opposite end of the spectrum from Python: "
    "static typing, manual memory management (mostly mitigated by RAII), and a build step "
    "before anything runs.",
    align="justify",
)
P(
    "The C++ side is organised around an object-oriented core. There is a Repository class for "
    "each entity (UsersRepo, StudentsRepo, CoursesRepo, EnrollmentsRepo, AttendanceRepo), a "
    "Service class that dispatches incoming requests to the correct repository, and a Net layer "
    "built directly on Winsock. Build and run with:",
    align="justify",
)
code_block("> g++ -std=c++17 -O2 -static -o server.exe server.cpp -lws2_32\n"
           "> g++ -std=c++17 -O2 -static -o client.exe client.cpp -lws2_32\n"
           "> server.exe          (terminal #1)\n"
           "> client.exe          (terminal #2)")

H("A. Console Login and Main Menu", level=2)
P(
    "On startup the C++ client tries to reach the server at 127.0.0.1:5066. If the connection "
    "fails, it prints a clear error and quits — better than hanging silently. After login the "
    "main menu is a familiar numbered list. Anything other than 0–4 is rejected on the spot, so "
    "you never end up in a state the program doesn’t expect.",
    align="justify",
)
figure("cpp_01_login.png", "Figure 11. C++ client — login prompt.")
figure("cpp_02_main_menu.png", "Figure 12. C++ client — main menu.")

H("B. Students Management", level=2)
P(
    "The Students menu mirrors the GUI: view all, search, add, update, delete. The view-all "
    "produces a properly aligned ASCII table — column widths are recomputed from the longest "
    "entry in each column, which keeps things readable even when names are unusually long:",
    align="justify",
)
figure("cpp_03_students_view_all.png", "Figure 13. C++ client — Students view-all.")
P(
    "Validation lives entirely on the server, which means the C++ client doesn’t need to "
    "duplicate any of the checks the Python client already does — it just forwards what the "
    "user typed and shows whatever the server says. Here’s what happens if you try to add a "
    "student with the ID \"abc\":",
    align="justify",
)
figure("cpp_04_students_add_error.png", "Figure 14. C++ client — Add student rejected (bad ID).")
P("And the happy path:")
figure("cpp_05_students_add_ok.png", "Figure 15. C++ client — Add student success.")

H("C. Courses Management", level=2)
P("Courses behave like students; the menu is symmetrical.")
figure("cpp_06_courses_view_all.png", "Figure 16. C++ client — Courses view-all.")

H("D. Enrollment", level=2)
P(
    "The Enrollment menu lets you list everyone’s enrollments, add a new one, or remove an "
    "existing one. When you list, the server joins the enrollment row with the student and "
    "course tables on its side so you see the names directly:",
    align="justify",
)
figure("cpp_07_enrollment_list.png", "Figure 17. C++ client — Enrollment list.")
figure("cpp_08_enrollment_add.png", "Figure 18. C++ client — Enrolling a student.")

H("E. Attendance", level=2)
P(
    "Attendance in the C++ client offers the same three actions as the Python tab — mark, "
    "report, summary — wrapped in a console menu. The report view filters by course code "
    "(leave blank to see everything):",
    align="justify",
)
figure("cpp_09_attendance_report.png", "Figure 19. C++ client — Attendance report for CS516.")
figure("cpp_10_attendance_summary.png", "Figure 20. C++ client — Attendance summary for CS516.")
P(
    "And here is the error case for an invalid date, which the server politely rejects so the "
    "client never persists garbage:",
    align="justify",
)
figure("cpp_11_attendance_error.png", "Figure 21. C++ client — Bad date format rejected.")

doc.add_page_break()

# =================================================================
#                       9. COMPARISON
# =================================================================
H("9. Comparison Between the Two Programming Languages", level=1)
P(
    "Now that both versions are working we can compare them honestly. We split the comparison "
    "into two parts as the rubric asks: a feature-by-feature table (syntax, dependencies, "
    "paradigm, memory management), followed by four side-by-side function comparisons.",
    align="justify",
)

H("Part 1 — Syntax, Dependencies, Paradigm, Memory Management", level=2)

cmp = doc.add_table(rows=5, cols=3)
cmp.style = "Light Grid Accent 1"
header = cmp.rows[0].cells
add_run(header[0].paragraphs[0], "Feature", bold=True)
add_run(header[1].paragraphs[0], "Python", bold=True)
add_run(header[2].paragraphs[0], "C++", bold=True)

# Syntax row
r = cmp.rows[1].cells
r[0].paragraphs[0].text = "Syntax"
syntax_py = (
    "Indentation defines blocks; no braces or semicolons.\n"
    "Print:   print(\"Hello\")\n"
    "Read:    name = input(\"name: \")\n"
    "If/else: if x > y: ...  else: ...\n"
    "Loop:    for i in range(n): ...\n"
    "Func:    def add(a, b): return a + b\n"
    "Variables don’t need a type."
)
syntax_cpp = (
    "Braces delimit blocks; every statement ends with ;\n"
    "Print:   std::cout << \"Hello\";\n"
    "Read:    std::cin >> name;\n"
    "If/else: if (x > y) { ... } else { ... }\n"
    "Loop:    for (int i = 0; i < n; ++i) { ... }\n"
    "Func:    int add(int a, int b) { return a + b; }\n"
    "Every variable has an explicit static type."
)
r[1].text = syntax_py
r[2].text = syntax_cpp

# Dependencies row
r = cmp.rows[2].cells
r[0].text = "Dependencies"
r[1].text = (
    "None outside the standard library. tkinter, sqlite3, socket and threading all ship with "
    "CPython, so running `python server.py` on any machine with Python 3.10+ just works. "
    "Cross-platform out of the box."
)
r[2].text = (
    "Requires a C++17 compiler and the Winsock library on Windows (or BSD sockets on Linux). "
    "We build with MinGW-w64 g++ 14.2 using `-lws2_32`. Static linking lets us ship a single "
    "self-contained .exe without any DLLs alongside it."
)

# Paradigm row
r = cmp.rows[3].cells
r[0].text = "Paradigm"
r[1].text = (
    "Multi-paradigm — procedural by default, with first-class support for OOP (classes, "
    "inheritance, polymorphism), functional features (lambdas, list/dict comprehensions, "
    "higher-order functions), and dynamic duck typing. Most of our Python code reads as "
    "object-oriented (each tab is a class)."
)
r[2].text = (
    "Mainly object-oriented and procedural; also supports generic programming through templates "
    "and a fair amount of functional style via lambdas and the STL algorithms. Our C++ server is "
    "purely OOP — a Service class owns one Repository per entity."
)

# Memory row
r = cmp.rows[4].cells
r[0].text = "Memory Management"
r[1].text = (
    "Automatic. CPython uses reference counting plus a cyclic garbage collector, so the "
    "developer almost never thinks about lifetime. The flip side is that you don’t get to "
    "control exactly when an object dies, which sometimes matters."
)
r[2].text = (
    "Mostly manual but rescued by RAII (Resource Acquisition Is Initialization). We use "
    "std::string and std::vector everywhere instead of raw new/delete, so memory is freed "
    "deterministically when objects go out of scope. No garbage collector — which is good for "
    "predictability but unforgiving if you forget."
)
P("Table 1. Comparison between Python and C++ along the four rubric axes.", italic=True, size=9)

# Part 2 — side by side functions
H("Part 2 — Side-by-Side Function Comparison", level=2)
P(
    "Below we walk through four of the system’s functions and show how each language expresses "
    "them. The Python side uses sqlite3 and JSON; the C++ side uses flat files and a small "
    "line-based protocol. We have not cherry-picked — these are the actual lines that ship in "
    "the project, just trimmed of imports.",
    align="justify",
)

# Function 1: Add Student
H("① Add Student", level=3)
P(
    "The first thing both implementations do is validate the incoming fields. Python does it "
    "with short helper functions that return True/False; C++ does the same thing with free "
    "functions that take strings. After validation, Python performs a parameterised INSERT "
    "against SQLite, while C++ asks its StudentsRepo to append a TAB-separated row to "
    "students.txt.",
    align="justify",
)
P("Python (server.py — handle_add_student):", bold=True)
code_block(
    "def handle_add_student(p):\n"
    "    if not is_valid_id(p.get(\"student_id\", \"\")):\n"
    "        return {\"ok\": False, \"error\": \"Student ID must be 10 digits\"}\n"
    "    if not p.get(\"name\", \"\").strip():\n"
    "        return {\"ok\": False, \"error\": \"Name is required\"}\n"
    "    if not is_valid_email(p.get(\"email\", \"\")):\n"
    "        return {\"ok\": False, \"error\": \"Invalid email\"}\n"
    "    if not is_valid_phone(p.get(\"phone\", \"\")):\n"
    "        return {\"ok\": False, \"error\": \"Phone must be 8 to 15 digits\"}\n"
    "    try:\n"
    "        level = int(p.get(\"level\"))\n"
    "        if level < 1 or level > 10: raise ValueError\n"
    "    except (TypeError, ValueError):\n"
    "        return {\"ok\": False, \"error\": \"Level must be a number from 1 to 10\"}\n"
    "    conn = get_conn()\n"
    "    try:\n"
    "        conn.execute(\"INSERT INTO students(student_id,name,email,phone,major,level) \"\n"
    "                     \"VALUES(?,?,?,?,?,?)\",\n"
    "                     (p['student_id'], p['name'].strip(), p['email'],\n"
    "                      p['phone'], p['major'].strip(), level))\n"
    "        conn.commit()\n"
    "    except sqlite3.IntegrityError:\n"
    "        return {\"ok\": False, \"error\": \"Student ID already exists\"}\n"
    "    return {\"ok\": True}"
)
P("C++ (server.cpp — Service::addStudent):", bold=True)
code_block(
    "Response addStudent(const Request& q) {\n"
    "    string sid = q.get(\"student_id\"), name  = trimCopy(q.get(\"name\"));\n"
    "    string email = q.get(\"email\"),    phone = q.get(\"phone\");\n"
    "    string major = trimCopy(q.get(\"major\")), level = q.get(\"level\");\n"
    "    if (!validId(sid))      return err(\"Student ID must be 10 digits\");\n"
    "    if (name.empty())       return err(\"Name is required\");\n"
    "    if (!validEmail(email)) return err(\"Invalid email\");\n"
    "    if (!validPhone(phone)) return err(\"Phone must be 8 to 15 digits\");\n"
    "    if (major.empty())      return err(\"Major is required\");\n"
    "    if (!isDigits(level))   return err(\"Level must be a number from 1 to 10\");\n"
    "    int lv = std::stoi(level);\n"
    "    if (lv < 1 || lv > 10)  return err(\"Level must be a number from 1 to 10\");\n"
    "    auto e = students_.add({sid, name, email, phone, major, level});\n"
    "    if (!e.empty()) return err(e);\n"
    "    return ok();\n"
    "}"
)
P(
    "What stands out: Python relies on duck-typed dictionaries (p.get(\"name\", \"\")) and a try/"
    "except around the database call to catch duplicates. C++ uses explicit types throughout and "
    "shifts the duplicate check inside the repository, which returns a non-empty string if "
    "something went wrong. Both versions hit roughly the same line count, but the C++ one has "
    "more visible structure.",
    align="justify",
)

# Function 2: View All Students
H("② View All Students", level=3)
P("Python issues one SELECT and lets sqlite3 return rows it can serialise straight to JSON.")
code_block(
    "def handle_list_students(p):\n"
    "    conn = get_conn()\n"
    "    rows = conn.execute(\n"
    "        \"SELECT student_id,name,email,phone,major,level FROM students \"\n"
    "        \"ORDER BY student_id\"\n"
    "    ).fetchall()\n"
    "    conn.close()\n"
    "    return {\"ok\": True, \"data\": [dict(r) for r in rows]}"
)
P("C++ reads its students.txt file and builds the rows by hand:")
code_block(
    "Response listStudents() {\n"
    "    Response r;\n"
    "    for (auto& row : students_.all()) {\n"
    "        if (row.size() < 6) continue;\n"
    "        r.rows.push_back({\n"
    "            {\"student_id\", row[0]}, {\"name\",  row[1]},\n"
    "            {\"email\",      row[2]}, {\"phone\", row[3]},\n"
    "            {\"major\",      row[4]}, {\"level\", row[5]}\n"
    "        });\n"
    "    }\n"
    "    return r;\n"
    "}"
)
P(
    "Here the comparison really tilts: SQLite gives you ordering, indexing, joins and "
    "transactions essentially for free. In the C++ version we trade that off for not having "
    "to install anything; the cost is that anything beyond a linear scan would need code we "
    "haven’t written.",
    align="justify",
)

# Function 3: Delete Student
H("③ Delete Student", level=3)
P("Python uses cascading deletes across three tables:")
code_block(
    "def handle_delete_student(p):\n"
    "    sid = p.get(\"student_id\", \"\")\n"
    "    conn = get_conn()\n"
    "    cur = conn.execute(\"DELETE FROM students WHERE student_id=?\", (sid,))\n"
    "    conn.execute(\"DELETE FROM enrollments WHERE student_id=?\", (sid,))\n"
    "    conn.execute(\"DELETE FROM attendance WHERE student_id=?\", (sid,))\n"
    "    conn.commit()\n"
    "    affected = cur.rowcount\n"
    "    conn.close()\n"
    "    if affected == 0:\n"
    "        return {\"ok\": False, \"error\": \"Student not found\"}\n"
    "    return {\"ok\": True}"
)
P("C++ relies on std::remove_if to filter the file’s rows in-memory and write them back:")
code_block(
    "Response deleteStudent(const Request& q) {\n"
    "    string sid = q.get(\"student_id\");\n"
    "    auto e = students_.remove(sid);\n"
    "    if (!e.empty()) return err(e);\n"
    "    enrollments_.removeAllForStudent(sid);\n"
    "    attendance_.removeAllForStudent(sid);\n"
    "    return ok();\n"
    "}\n\n"
    "// inside StudentsRepo:\n"
    "string remove(const string& sid) {\n"
    "    auto rows = store_.readAll();\n"
    "    size_t before = rows.size();\n"
    "    rows.erase(std::remove_if(rows.begin(), rows.end(),\n"
    "        [&](const vector<string>& r){ return !r.empty() && r[0] == sid; }),\n"
    "        rows.end());\n"
    "    if (rows.size() == before) return \"Student not found\";\n"
    "    store_.writeAll(rows);\n"
    "    return \"\";\n"
    "}"
)
P(
    "Python’s version is shorter and looks almost like pseudo-code. The C++ version reads more "
    "explicit because of the iterators and the lambda — but in return we get a compiler that "
    "would catch us if we tried to use rows after std::remove_if invalidated the iterators.",
    align="justify",
)

# Function 4: Attendance summary
H("④ Per-course Attendance Summary", level=3)
P(
    "This one shows the biggest gap. Python pushes the entire computation into SQL with a "
    "GROUP BY:",
    align="justify",
)
code_block(
    "def handle_attendance_summary(p):\n"
    "    code = p.get(\"course_code\", \"\").strip().upper()\n"
    "    conn = get_conn()\n"
    "    rows = conn.execute(\n"
    "        \"SELECT a.student_id, s.name, \"\n"
    "        \"  SUM(CASE WHEN a.status='Present' THEN 1 ELSE 0 END) AS present_count, \"\n"
    "        \"  SUM(CASE WHEN a.status='Absent'  THEN 1 ELSE 0 END) AS absent_count \"\n"
    "        \"FROM attendance a LEFT JOIN students s ON s.student_id=a.student_id \"\n"
    "        \"WHERE a.course_code=? GROUP BY a.student_id ORDER BY s.name\",\n"
    "        (code,)\n"
    "    ).fetchall()\n"
    "    conn.close()\n"
    "    return {\"ok\": True, \"data\": [dict(r) for r in rows]}"
)
P("C++ has to roll its own aggregation:")
code_block(
    "Response attendanceSummary(const Request& q) {\n"
    "    Response r;\n"
    "    string code = toUpper(trimCopy(q.get(\"course_code\")));\n"
    "    map<string, std::pair<int,int>> counts;  // sid -> {present, absent}\n"
    "    for (auto& row : attendance_.all()) {\n"
    "        if (row.size() < 4) continue;\n"
    "        if (!code.empty() && row[1] != code) continue;\n"
    "        if (row[3] == \"Present\") counts[row[0]].first++;\n"
    "        else                       counts[row[0]].second++;\n"
    "    }\n"
    "    map<string, string> sname;\n"
    "    for (auto& s : students_.all()) if (s.size() >= 2) sname[s[0]] = s[1];\n"
    "    for (auto& kv : counts) {\n"
    "        r.rows.push_back({\n"
    "            {\"student_id\", kv.first}, {\"name\", sname[kv.first]},\n"
    "            {\"present_count\", std::to_string(kv.second.first)},\n"
    "            {\"absent_count\",  std::to_string(kv.second.second)}\n"
    "        });\n"
    "    }\n"
    "    return r;\n"
    "}"
)
P(
    "Python expresses the whole thing in one SQL statement that the database engine optimises; "
    "C++ takes more code but is also more transparent about what is happening. For a small "
    "department both run in essentially zero time, but on bigger data the SQL approach scales "
    "better without us having to touch the code.",
    align="justify",
)

# Conclusion of comparison
H("Verdict in one paragraph", level=2)
P(
    "Python was faster to write and the GUI gave us a polished UI for almost no effort. C++ "
    "forced us to be more explicit and produced a smaller, dependency-free binary, but every "
    "feature took noticeably longer to land. For a project of this scope Python is the more "
    "productive choice; for a project where performance or distribution size matter, C++ would "
    "pay back the upfront cost. They are not in competition — they answer different questions.",
    align="justify",
)

doc.add_page_break()

# =================================================================
#                        10. ATTACHMENTS
# =================================================================
H("10. Attachments", level=1)
P(
    "The full source code accompanies this report in the project archive. The folder layout is:",
    align="justify",
)
code_block(
    "CS516_Project/\n"
    "├── Python/\n"
    "│   ├── server.py        # TCP server + sqlite3 backend\n"
    "│   ├── client.py        # tkinter GUI client\n"
    "│   └── smartsms.db      # created on first run\n"
    "├── Cpp/\n"
    "│   ├── common.h         # shared protocol types\n"
    "│   ├── server.cpp       # OOP server + flat-file storage\n"
    "│   ├── client.cpp       # console client\n"
    "│   ├── server.exe       # pre-built (MinGW-w64 g++ 14.2)\n"
    "│   ├── client.exe       # pre-built\n"
    "│   └── data/            # created on first run\n"
    "├── Screenshots/         # every figure used in this report\n"
    "└── Report/\n"
    "    └── CS516_Project_Report.docx"
)
P("How to run:", bold=True)
P("Python side:")
code_block(
    "$ cd CS516_Project/Python\n"
    "$ python server.py        # terminal 1\n"
    "$ python client.py        # terminal 2\n"
    "# default login: admin / admin123"
)
P("C++ side:")
code_block(
    "> cd CS516_Project\\Cpp\n"
    "> server.exe              :: terminal 1\n"
    "> client.exe              :: terminal 2\n"
    ":: default login: admin / admin123\n\n"
    ":: To rebuild from source (Windows + MinGW-w64):\n"
    "> g++ -std=c++17 -O2 -static -o server.exe server.cpp -lws2_32\n"
    "> g++ -std=c++17 -O2 -static -o client.exe client.cpp -lws2_32"
)

# =================================================================
#                       11. REFERENCES
# =================================================================
H("11. References", level=1)
for r in [
    "Python Software Foundation. \"sqlite3 — DB-API 2.0 interface for SQLite databases\". "
    "Python 3 documentation.",
    "Python Software Foundation. \"tkinter — Python interface to Tcl/Tk\". Python 3 documentation.",
    "ISO/IEC 14882:2017. Programming languages — C++. ISO, 2017.",
    "Microsoft. \"Winsock 2 (Windows Sockets 2) — Reference\". Microsoft Learn, accessed 2026.",
    "B. Stroustrup, A Tour of C++, 2nd ed., Addison-Wesley, 2018.",
    "Course notes and lecture slides, CS516 Advanced Programming Language, IAU, 2025–2026.",
]:
    bullet(r)

# Save
doc.save(OUT)
print("Saved:", OUT)
print("Size:", os.path.getsize(OUT), "bytes")
