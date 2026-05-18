"""
Build a one-page Word guide explaining how to test/run the project.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "CS516_Test_Guide.docx")

NAVY = "1F3A5F"
GOLD = "E3B341"
ACCENT = "1F6FEB"
GREEN = "2EA043"

doc = Document()

# Tight page setup so it fits on one A4 page
for section in doc.sections:
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(2)


def run(p, text, *, bold=False, italic=False, size=None, color=None, font=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if font:
        r.font.name = font
    return r


def shade(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def section_bar(label, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    shade(p, color)
    run(p, " " + label, bold=True, size=11, color="FFFFFF")


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_after = Pt(2)
    shade(p, "F4F6F8")
    run(p, text, font="Consolas", size=9)


def bullet(text, bold_part=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    if bold_part:
        run(p, bold_part, bold=True, size=10)
        run(p, " " + text, size=10)
    else:
        run(p, text, size=10)


# ---------------- Header ----------------
hdr = doc.add_paragraph()
hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
run(hdr, "CS 516 — Test Guide\n", bold=True, size=18, color=NAVY)
run(hdr, "Smart Student Management & Attendance System  ·  How to run and test the project",
    italic=True, size=10, color="5F6B7C")

# subtle gold separator
sep = doc.add_paragraph()
shade(sep, GOLD)
sep.paragraph_format.space_after = Pt(4)
run(sep, " ", size=2)

# ---------------- Prerequisites ----------------
section_bar("Prerequisites", NAVY)
bullet("Python 3.10+ (tkinter, sqlite3, socket all bundled).  Git is needed once to clone the repo.", "You need:")
bullet("admin / admin123 (full access)   ·   ahmed / ahmed123   ·   abdullah / abdullah123 (instructor)", "Default logins:")

# ---------------- Download ----------------
section_bar("Step 1 — Download the project", "8C1D40")
p = doc.add_paragraph()
run(p, "Open a PowerShell window in any folder you like and clone the repository:", size=10)
code("PS> git clone https://github.com/abdullahalsaedi2003-dotcom/adaved-programing-project.git\nPS> cd adaved-programing-project")
p = doc.add_paragraph()
run(p, "No Git? ", bold=True, size=10)
run(p, "Open the repo in your browser, click ", size=10)
run(p, "Code → Download ZIP", bold=True, size=10)
run(p, ", then unzip it. The pre-built ", size=10)
run(p, "server.exe", font="Consolas", size=10)
run(p, " and ", size=10)
run(p, "client.exe", font="Consolas", size=10)
run(p, " are included — no compilation needed.", size=10)

# ---------------- Python ----------------
section_bar("Step 2 — Test the Python Version  (Tkinter GUI · SQLite · port 5055)", ACCENT)
p = doc.add_paragraph()
run(p, "Open ", size=10); run(p, "two", bold=True, size=10)
run(p, " PowerShell windows. Server in the first, GUI client in the second:", size=10)
code("PS> cd Python\nPS> python server.py         :: terminal 1 — \"Smart SMS server listening on 127.0.0.1:5055\"\nPS> python client.py         :: terminal 2 — login window appears (admin / admin123 pre-filled)")

# ---------------- C++ ----------------
section_bar("Step 3 — Test the C++ Version  (Console · flat files · port 5066)", GREEN)
p = doc.add_paragraph()
run(p, "Same idea, two PowerShell windows:", size=10)
code("PS> cd Cpp\nPS> .\\server.exe             :: terminal 1 — \"Smart SMS (C++) server listening on 127.0.0.1:5066\"\nPS> .\\client.exe             :: terminal 2 — press Enter twice to accept default admin login")

# ---------------- Test scenarios ----------------
section_bar("Step 4 — Test Scenarios (try in either client)", NAVY)
bullet("View All → confirm the 3 seeded students appear. Add ID 2220001111 (try \"abc\" first — rejected).", "Students:")
bullet("CS516, CS411, CS324 pre-loaded. Add CS401 / Database Systems / 3 / Dr. Mona.", "Courses:")
bullet("Enroll 2220007240 in CS516 / 2025-2. Try the same combo twice — server blocks duplicates.", "Enroll:")
bullet("Mark Present for 2220007240 / CS516 / today. Try date 14-5-2026 — server rejects bad format.", "Attendance:")
bullet("Switch to Summary, enter CS516, see per-student present/absent counts.", "Summary:")

# ---------------- Troubleshooting ----------------
section_bar("Troubleshooting", "8C1D40")
bullet("Server not running, or port busy. Close old terminals, or:  Get-NetTCPConnection -LocalPort 5055,5066 | ForEach-Object { Stop-Process -Id $_.OwningProcess -Force }", "“Cannot reach the server”:")
bullet("Delete  Python\\smartsms.db  and  Cpp\\data\\  — both are regenerated with seed data on next start.", "Clean reset:")

# ---------------- Footer ----------------
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.paragraph_format.space_before = Pt(4)
run(foot,
    "Source code & full report  ·  github.com/abdullahalsaedi2003-dotcom/adaved-programing-project",
    italic=True, size=9, color="5F6B7C")

doc.save(OUT)
print("Saved:", OUT, os.path.getsize(OUT), "bytes")
